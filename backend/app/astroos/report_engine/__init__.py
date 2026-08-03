"""Report Engine — PDF / Word / HTML / Markdown / JSON + multi-language."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

SUPPORTED_LANGS = ("en", "mr", "hi", "gu", "kn", "ta", "te")

LANG_LABELS = {
    "en": "English",
    "mr": "Marathi",
    "hi": "Hindi",
    "gu": "Gujarati",
    "kn": "Kannada",
    "ta": "Tamil",
    "te": "Telugu",
}


_LANG_FIELD = {
    "en": "english",
    "mr": "marathi",
    "hi": "hindi",
    "gu": "gujarati",
    "kn": "kannada",
    "ta": "tamil",
    "te": "telugu",
}


def _source_text(src: dict[str, Any], language: str) -> str:
    field = _LANG_FIELD.get(language, language)
    return (
        src.get(field)
        or src.get(language)
        or src.get("english")
        or src.get("en")
        or ""
    )


def to_markdown(report: dict[str, Any], language: str = "en") -> str:
    lines = [
        f"# {report.get('title', 'AstroOS Report')}",
        f"Language: {LANG_LABELS.get(language, language)}",
        "",
        report.get("disclaimer", ""),
        "",
        "## Conclusions",
        "",
    ]
    for c in report.get("conclusions", []):
        lines.append(f"### {c['title']} ({c['confidence']}%)")
        lines.append(c.get("summary", ""))
        lines.append("")
        lines.append("**Evidence**")
        for e in c.get("evidence", []):
            lines.append(f"- {e}")
        lines.append("")
        lines.append("**Rules:** " + ", ".join(c.get("used_rules", [])))
        lines.append("")
        for s in c.get("sources", []):
            translated = _source_text(s, language)
            lines.append(
                f"- Source: {s.get('text')} / {s.get('sloka_id')} — {translated}"
            )
        lines.append("")
        if c.get("ai_explanation"):
            lines.append("**AI Explanation**")
            lines.append(c["ai_explanation"])
            lines.append("")
    return "\n".join(lines)


def to_html(report: dict[str, Any], language: str = "en") -> str:
    body = to_markdown(report, language).replace("\n", "<br/>\n")
    return (
        "<!doctype html><html><head><meta charset='utf-8'></head>"
        f"<body style='font-family:Georgia,serif;padding:2rem'>{body}</body></html>"
    )


def to_json(report: dict[str, Any]) -> str:
    return json.dumps(report, ensure_ascii=False, indent=2)


def to_pdf(report: dict[str, Any], out_path: Path, language: str = "en") -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(out_path), pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, report.get("title", "AstroOS Report")[:80])
    y -= 30
    c.setFont("Helvetica", 9)
    for line in to_markdown(report, language).splitlines():
        if y < 60:
            c.showPage()
            y = height - 50
            c.setFont("Helvetica", 9)
        # Helvetica can't render all Indic; keep ASCII-safe slice for PDF engine
        safe = line.encode("ascii", "replace").decode("ascii")
        c.drawString(40, y, safe[:110])
        y -= 12
    c.save()
    return out_path


def to_docx(report: dict[str, Any], out_path: Path, language: str = "en") -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx import Document
    except ImportError:
        # Fallback: write rich-text-ish markdown as .docx.txt sibling note
        fallback = out_path.with_suffix(".docx.md")
        fallback.write_text(to_markdown(report, language), encoding="utf-8")
        return fallback

    doc = Document()
    doc.add_heading(report.get("title", "AstroOS Report"), level=1)
    doc.add_paragraph(f"Language: {LANG_LABELS.get(language, language)}")
    doc.add_paragraph(report.get("disclaimer", ""))
    for c in report.get("conclusions", []):
        doc.add_heading(f"{c['title']} ({c['confidence']}%)", level=2)
        doc.add_paragraph(c.get("summary", ""))
        doc.add_paragraph("Evidence:")
        for e in c.get("evidence", []):
            doc.add_paragraph(e, style="List Bullet")
        doc.add_paragraph("Rules: " + ", ".join(c.get("used_rules", [])))
        for s in c.get("sources", []):
            doc.add_paragraph(
                f"{s.get('text')} / {s.get('sloka_id')}: {_source_text(s, language)}"
            )
        if c.get("ai_explanation"):
            doc.add_paragraph("AI Explanation:")
            doc.add_paragraph(c["ai_explanation"])
    doc.save(str(out_path))
    return out_path


def generate_report_bundle(
    report: dict[str, Any],
    stem: str = "astroos_report",
    language: str = "en",
) -> dict[str, str]:
    import os
    import tempfile

    language = language if language in SUPPORTED_LANGS else "en"
    preferred = Path(__file__).resolve().parents[2] / "storage" / "reports"
    # Vercel / serverless filesystems are read-only except /tmp
    if os.environ.get("VERCEL") or os.environ.get("AWS_LAMBDA_FUNCTION_NAME"):
        storage = Path(tempfile.gettempdir()) / "astroos-reports"
    else:
        storage = preferred
    try:
        storage.mkdir(parents=True, exist_ok=True)
        probe = storage / ".write_probe"
        probe.write_text("ok", encoding="utf-8")
        probe.unlink(missing_ok=True)
    except OSError:
        storage = Path(tempfile.gettempdir()) / "astroos-reports"
        storage.mkdir(parents=True, exist_ok=True)

    md_path = storage / f"{stem}.md"
    html_path = storage / f"{stem}.html"
    json_path = storage / f"{stem}.json"
    pdf_path = storage / f"{stem}.pdf"
    docx_path = storage / f"{stem}.docx"
    md_path.write_text(to_markdown(report, language), encoding="utf-8")
    html_path.write_text(to_html(report, language), encoding="utf-8")
    json_path.write_text(to_json(report), encoding="utf-8")
    to_pdf(report, pdf_path, language)
    to_docx(report, docx_path, language)
    return {
        "markdown": str(md_path),
        "html": str(html_path),
        "json": str(json_path),
        "pdf": str(pdf_path),
        "docx": str(docx_path),
        "language": language,
        "supported_languages": list(SUPPORTED_LANGS),
    }
